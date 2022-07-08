import docx
from docx.shared import Pt


doc = docx.Document()


doc.add_heading('Dados da Disciplina : ', 0)


dis = input(
   'Qual é a disciplina ? ')     
para = doc.add_paragraph()
p = para.add_run(
   '-Disciplina : ')
p.font.size = Pt(15)


user_input = para.add_run(
   f'{dis}.')
user_input.font.size = Pt(15)
user_input.bold = True


time = input(
   'Qual é a carga horária ? ')
para = doc.add_paragraph(style = 'List Bullet')
bold_para = para.add_run(
   'Com carga horária de ')
bold_para.font.size = Pt(15)


user_input = para.add_run(
   f'{time} horas.')
user_input.font.size = Pt(15)
user_input.bold = True


para = doc.add_paragraph()
bold_para=para.add_run(
   'Conteúdo da disciplina : ')
bold_para.font.size = Pt(15)
bold_para.bold = True


tab = doc.add_table(rows=1, cols=2)
tab.style="Medium Grid 1 Accent 4"
cels = tab.rows[0].cells
cels[0].text = 'Número:'
cels[0].width = 5
cels[1].text = "Assunto:"


for num in range(1, 5):
   assunto = input(
      f'Qual é o assunto número {num} ? ')
   dados = tab.add_row().cells
   dados[0].text = str(num)
   dados[0].width = 5
   dados[1].text = assunto


doc.save('/home/@/Documents/whatever/meu_Docx.docx')
